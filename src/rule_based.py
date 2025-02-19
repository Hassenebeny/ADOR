import logging
import re
import docx

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def get_next_nonempty_paragraph(paragraphs, start_index):
    """
    Retourne le texte du prochain paragraphe non vide après start_index.
    """
    for j in range(start_index + 1, len(paragraphs)):
        p_text = paragraphs[j].text.strip()
        if p_text:
            return p_text
    return ""

def extract_entities_from_table(doc):
    """
    Parcourt tous les tableaux du document DOCX et extrait les paires clé/valeur selon un mapping.
    """
    data = {
        "Counterparty": "",
        "Initial Valuation Date": "",
        "Notional": "",
        "Valuation Date": "",
        "Maturity": "",
        "Underlying": "",
        "Coupon": "",
        "Barrier": "",
        "Calendar": ""
    }
    mapping = {
        "party a": "Counterparty",
        "initial valuation date": "Initial Valuation Date",
        "notional amount": "Notional",
        "valuation date": "Valuation Date",
        "termination date": "Maturity",
        "underlying": "Underlying",
        "coupon": "Coupon",
        "barrier": "Barrier",
        "business day": "Calendar"
    }
    
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            key_raw = cells[0].text.strip().lower()
            value_raw = cells[1].text.strip()
            for known_key, entity_name in mapping.items():
                if known_key in key_raw:
                    if entity_name == "Barrier":
                        match = re.search(r"(\d{1,3}(?:\.\d+)?%|\d{1,3}%)", value_raw)
                        data[entity_name] = match.group(1) if match else value_raw
                    else:
                        data[entity_name] = value_raw
                    break
    return data

def extract_entities_from_paragraphs(doc):
    """
    Extraction rule-based depuis les paragraphes si aucun tableau n'est présent.
    Pour chaque libellé détecté, la valeur correspondante est prise dans le paragraphe suivant non vide.
    """
    data = {
        "Counterparty": "",
        "Initial Valuation Date": "",
        "Notional": "",
        "Valuation Date": "",
        "Maturity": "",
        "Underlying": "",
        "Coupon": "",
        "Barrier": "",
        "Calendar": ""
    }
    mapping_keys = {
        "party a": "Counterparty",
        "initial valuation date": "Initial Valuation Date",
        "notional amount": "Notional",
        "valuation date": "Valuation Date",
        "termination date": "Maturity",
        "underlying": "Underlying",
        "coupon": "Coupon",
        "barrier": "Barrier",
        "business day": "Calendar"
    }
    
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        text_lower = p.text.strip().lower()
        for key, entity in mapping_keys.items():
            if key in text_lower:
                value = get_next_nonempty_paragraph(paragraphs, i)
                if entity == "Barrier":
                    match = re.search(r"(\d{1,3}(?:\.\d+)?%|\d{1,3}%)", value)
                    value = match.group(1) if match else value
                data[entity] = value
    return data

def rule_based_extraction(file_path):
    """
    Ouvre le document DOCX et détecte si celui-ci contient des tableaux.
    Utilise l'extraction table-based si des tableaux sont détectés, sinon utilise l'extraction par paragraphes.
    Retourne un dictionnaire contenant les entités extraites.
    """
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logger.error(f"Erreur lors de l'ouverture du fichier DOCX : {e}")
        return {}

    if len(doc.tables) > 0:
        logger.info("Tableaux détectés dans le document DOCX. Utilisation de l'extraction table-based.")
        data_extracted = extract_entities_from_table(doc)
    else:
        logger.info("Aucun tableau détecté. Utilisation de l'extraction par paragraphes.")
        data_extracted = extract_entities_from_paragraphs(doc)

    return data_extracted
