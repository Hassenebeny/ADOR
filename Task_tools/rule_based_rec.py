import re
from docx import Document
import json

def extraire_texte_docx(docx_path):
    """
    Extrait le texte du document DOCX, en parcourant les paragraphes et les tableaux.
    Retourne une liste de chaînes de caractères.
    """
    doc = Document(docx_path)
    contenu = []

    # Extraction depuis les paragraphes
    for para in doc.paragraphs:
        ligne = para.text.strip()
        if ligne:
            contenu.append(ligne)
    
    # Extraction depuis les tableaux
    for table in doc.tables:
        for row in table.rows:
            # On concatène les textes de toutes les cellules de la ligne
            ligne = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if ligne:
                contenu.append(ligne)
    
    return contenu

def extraire_entites(docx_path):
    """
    Lit le document DOCX et extrait des entités nommées à l'aide d'un parseur à base de règles..
    """
    contenu = extraire_texte_docx(docx_path)
    texte = "\n".join(contenu)
    
    # Dictionnaire qui contiendra les entités extraites
    entites = {
        "Counterparty": [],
        "Initial Valuation Date": None,
        "Notional": None,
        "Valuation Date": None,
        "Maturity": None,
        "Underlying": None,
        "Coupon": None,
        "Barrier": None,
        "Calendar": None
    }
    
    # Règles (expressions régulières) pour extraire les entités
    
    pattern_party_a = r"(Party A)\s*[:\n]?\s*(.+)"
    pattern_party_b = r"(Party B)\s*[:\n]?\s*(.+)"
    pattern_initial_valuation = r"(Initial Valuation Date)\s*[:\n]?\s*(.+)"
    pattern_notional = r"(Notional Amount.*?|Notional)\s*[:\n]?\s*(.+)"
    pattern_valuation = r"(Valuation Date)\s*[:\n]?\s*(.+)"
    pattern_maturity = r"(Termination Date|Maturity)\s*[:\n]?\s*(.+)"
    pattern_underlying = r"(Underlying)\s*[:\n]?\s*(.+)"
    pattern_coupon = r"(Coupon.*?|Coupon \(C\))\s*[:\n]?\s*(.+)"
    pattern_barrier = r"(Barrier.*?|Barrier \(B\))\s*[:\n]?\s*(.+)"
    pattern_calendar = r"(Business Day|Calendar)\s*[:\n]?\s*(.+)"
    
    # Fonction utilitaire pour extraire une valeur à partir d'un motif
    def extraire_valeur(pattern, texte):
        match = re.search(pattern, texte, re.IGNORECASE)
        if match:
            return match.group(2).strip()
        return None

    # Extraction des contreparties en parcourant chaque ligne
    for ligne in contenu:
        match_a = re.search(pattern_party_a, ligne, re.IGNORECASE)
        if match_a:
            entites["Counterparty"].append(match_a.group(2).strip())
        match_b = re.search(pattern_party_b, ligne, re.IGNORECASE)
        if match_b:
            entites["Counterparty"].append(match_b.group(2).strip())

    # Extraction des autres entités à partir du texte global
    entites["Initial Valuation Date"] = extraire_valeur(pattern_initial_valuation, texte)
    entites["Notional"] = extraire_valeur(pattern_notional, texte)
    entites["Valuation Date"] = extraire_valeur(pattern_valuation, texte)
    entites["Maturity"] = extraire_valeur(pattern_maturity, texte)
    entites["Underlying"] = extraire_valeur(pattern_underlying, texte)
    entites["Coupon"] = extraire_valeur(pattern_coupon, texte)
    entites["Barrier"] = extraire_valeur(pattern_barrier, texte)
    entites["Calendar"] = extraire_valeur(pattern_calendar, texte)
    
    return entites

def generer_output_json(entites, output_path="resultat.json"):
    """
    Génère un fichier JSON à partir du dictionnaire d'entités extraites.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(entites, f, ensure_ascii=False, indent=4)
    print(f"Fichier de sortie généré: {output_path}")



if __name__ == "__main__":
    chemin_fichier_docx = "data/ZF4894_ALV_07Aug2026_physical.docx"
    resultat = extraire_entites(chemin_fichier_docx)
    
    print("Entités extraites :")
    for cle, valeur in resultat.items():
        print(f"{cle} : {valeur}")
    
    generer_output_json(resultat, "entites_extraites.json")
